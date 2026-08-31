"""Ingestion tests. No API key required — these only exercise loading and
chunking, not embeddings. Uses tests/fixtures/, not data/raw/, since data/raw
holds whatever real documents the project owner has actually ingested."""
import sqlite3
from pathlib import Path

import pandas as pd
from langchain_core.documents import Document

from src.ingestion.chunking import split_documents
from src.ingestion.loaders import (
    load_directory,
    load_prose_directory,
    strip_shared_boilerplate,
)
from src.ingestion.structured_loaders import load_sql_table

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"


def test_load_directory_finds_sample_docs():
    docs = load_directory(FIXTURES_DIR)
    sources = {doc.metadata["source"] for doc in docs}
    assert "sample.md" in sources


def test_load_directory_handles_structured_formats():
    docs = load_directory(FIXTURES_DIR)
    sources = {doc.metadata["source"] for doc in docs}
    assert "sample.csv" in sources
    assert "sample.json" in sources

    csv_rows = [doc for doc in docs if doc.metadata["source"] == "sample.csv"]
    assert len(csv_rows) == 3  # one Document per data row
    assert "TICK-001" in csv_rows[0].page_content
    assert "severity" in csv_rows[0].page_content.lower()

    json_rows = [doc for doc in docs if doc.metadata["source"] == "sample.json"]
    assert len(json_rows) == 2
    assert any("Sprocket" in doc.page_content for doc in json_rows)


def test_load_prose_directory_excludes_structured_formats():
    docs = load_prose_directory(FIXTURES_DIR)
    sources = {doc.metadata["source"] for doc in docs}
    assert sources == {"sample.md"}  # not sample.csv or sample.json


def test_load_directory_handles_excel(tmp_path):
    xlsx_path = tmp_path / "sample.xlsx"
    with pd.ExcelWriter(xlsx_path) as writer:
        pd.DataFrame({"a": [1, 2]}).to_excel(writer, sheet_name="sheet_one", index=False)
        pd.DataFrame({"b": [3, 4]}).to_excel(writer, sheet_name="sheet_two", index=False)

    docs = load_directory(tmp_path)
    sources = {doc.metadata["source"] for doc in docs}
    assert "sample.xlsx" in sources
    assert {doc.metadata["sheet"] for doc in docs} == {"sheet_one", "sheet_two"}


def test_load_directory_handles_docx(tmp_path):
    """Real bug, found by auditing the project for unused code: requirements.txt
    listed `python-docx`, which the app never actually imports — .docx loading
    goes through langchain's Docx2txtLoader, which needs the differently-named
    `docx2txt` package instead. That package was missing entirely, so dropping
    a real .docx into data/raw crashed with `ModuleNotFoundError: No module
    named 'docx2txt'` the moment anyone actually tried it. Undetected because
    no test exercised .docx loading at all. Fixed by swapping the dependency;
    this closes the coverage gap that let it happen. python-docx is only
    needed here, to author the fixture — the app's own runtime path never
    imports it."""
    from docx import Document as DocxDocument

    docx_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("Refunds are processed within 5 business days.")
    doc.save(docx_path)

    docs = load_directory(tmp_path)
    sources = {doc.metadata["source"] for doc in docs}
    assert "sample.docx" in sources
    assert any("5 business days" in doc.page_content for doc in docs)


def test_load_sql_table(tmp_path):
    db_path = tmp_path / "test.db"
    conn = sqlite3.connect(db_path)
    conn.execute("CREATE TABLE widgets (id INTEGER, name TEXT, price REAL)")
    conn.execute("INSERT INTO widgets VALUES (1, 'sprocket', 4.5), (2, 'gizmo', 9.0)")
    conn.commit()
    conn.close()

    docs = load_sql_table(f"sqlite:///{db_path.as_posix()}", table="widgets")
    assert len(docs) == 2
    assert "sprocket" in docs[0].page_content
    assert docs[0].metadata["table"] == "widgets"
    assert docs[0].metadata["source"] == "db:widgets"


def test_split_documents_respects_chunk_size():
    docs = load_directory(FIXTURES_DIR)
    chunks = split_documents(docs)
    assert len(chunks) >= len(docs)
    for chunk in chunks:
        assert len(chunk.page_content) <= 1000 + 200  # chunk_size + slack for separators
        assert "chunk_id" in chunk.metadata
        assert "source" in chunk.metadata


def _page(source: str, *lines: str) -> Document:
    return Document(page_content="\n".join(lines), metadata={"source": source})


def test_strip_shared_boilerplate_removes_lines_repeated_across_pages():
    """Real, measured problem this solves: 51% of all scraped web
    line-content in this project's own corpus was the same ~100 nav/footer
    lines duplicated verbatim across all 16 pages, directly costing
    retrieval precision."""
    nav = ["Home", "About", "Contact", "Apply Now"]
    docs = [
        _page("http://x/a", *nav, "Tuition is 5000 EUR per semester."),
        _page("http://x/b", *nav, "The library closes at 8pm."),
        _page("http://x/c", *nav, "Enrollment opens in March."),
        _page("http://x/d", *nav, "Parking permits cost 50 EUR."),
    ]

    cleaned = strip_shared_boilerplate(docs)

    assert len(cleaned) == 4
    for doc in cleaned:
        for nav_line in nav:
            assert nav_line not in doc.page_content
    assert "Tuition is 5000 EUR per semester." in cleaned[0].page_content
    assert "Parking permits cost 50 EUR." in cleaned[3].page_content


def test_strip_shared_boilerplate_keeps_content_unique_to_one_page():
    docs = [
        _page("http://x/a", "Shared nav", "Unique fact A"),
        _page("http://x/b", "Shared nav", "Unique fact B"),
        _page("http://x/c", "Shared nav", "Unique fact C"),
        _page("http://x/d", "Shared nav", "Unique fact D"),
    ]

    cleaned = strip_shared_boilerplate(docs)

    assert [d.page_content for d in cleaned] == [
        "Unique fact A",
        "Unique fact B",
        "Unique fact C",
        "Unique fact D",
    ]


def test_strip_shared_boilerplate_is_a_noop_for_too_few_pages():
    """"Appears on 1 of 2 pages" is 50% but means nothing — with almost no
    pages there's no reliable repetition signal, and guessing would risk
    deleting real content. Skipping is the safe default."""
    docs = [
        _page("http://x/a", "Same line", "Fact A"),
        _page("http://x/b", "Same line", "Fact B"),
    ]

    assert strip_shared_boilerplate(docs) == docs


def test_strip_shared_boilerplate_drops_a_page_that_was_entirely_boilerplate():
    docs = [
        _page("http://x/a", "Nav", "Real content A"),
        _page("http://x/b", "Nav", "Real content B"),
        _page("http://x/c", "Nav", "Real content C"),
        _page("http://x/nav-only", "Nav"),
    ]

    cleaned = strip_shared_boilerplate(docs)

    sources = [d.metadata["source"] for d in cleaned]
    assert "http://x/nav-only" not in sources  # nothing left worth indexing
    assert len(cleaned) == 3


def test_strip_shared_boilerplate_preserves_metadata():
    docs = [
        Document(page_content="Nav\nFact A", metadata={"source": "http://x/a", "extra": 1}),
        Document(page_content="Nav\nFact B", metadata={"source": "http://x/b", "extra": 2}),
        Document(page_content="Nav\nFact C", metadata={"source": "http://x/c", "extra": 3}),
        Document(page_content="Nav\nFact D", metadata={"source": "http://x/d", "extra": 4}),
    ]

    cleaned = strip_shared_boilerplate(docs)

    assert [d.metadata["extra"] for d in cleaned] == [1, 2, 3, 4]


def test_split_documents_assigns_deterministic_ids():
    """Real bug, found live against the project's own data store: chunks
    had no explicit Chroma id, so add_documents() generated a fresh random
    one every call — re-ingesting the same file duplicated it instead of
    upserting. 1340 of 1515 real chunks ended up exact duplicates after a
    routine re-ingest that only meant to add two new URLs. `chunk_id`
    (a running index within one split_documents() call) isn't stable
    across separate calls either, which is why the fix is a content hash,
    not that field."""
    docs = load_directory(FIXTURES_DIR)

    chunks_first_run = split_documents(docs)
    chunks_second_run = split_documents(docs)  # simulates re-ingesting the same files

    ids_first = [c.id for c in chunks_first_run]
    ids_second = [c.id for c in chunks_second_run]

    assert all(ids_first)  # every chunk actually got an id, not None
    assert ids_first == ids_second  # same (source, content) -> same ids, every time
    assert len(set(ids_first)) == len(ids_first)  # distinct chunks -> distinct ids
